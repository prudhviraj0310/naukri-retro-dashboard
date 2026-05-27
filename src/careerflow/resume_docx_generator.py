import os
import json
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

def build_tailored_resume_docx(
    candidate_knowledge: dict,
    tailored_data: dict,
    output_path: str = "resumes/tailored_resume.docx"
) -> str:
    """
    Generate an ATS-optimized, beautifully styled SDE resume in .docx format.
    Dynamically weaves in AI-tailored professional summaries and experience bullet points.
    """
    doc = Document()
    
    # Page setup - Standard 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Style definitions
    styles = doc.styles
    
    # Configure base font style (Calibri / Arial, clean and highly ATS compliant)
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Sleek off-black
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # Create Header Info
    identity = candidate_knowledge.get("identity", {})
    name = identity.get("name", "Prudhvi Raj")
    role = identity.get("role", "Backend & AI Automation Engineer")
    
    # Defaults
    defaults = candidate_knowledge.get("questionnaire_defaults", {})
    email = defaults.get("email", "prudhvi.raj.nakari@gmail.com")
    phone = defaults.get("phone", "8142498424")
    location = defaults.get("preferred_location", "Bangalore")
    
    # 1. Main Header Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name_run = p_name.add_run(name)
    p_name_run.font.size = Pt(20)
    p_name_run.font.bold = True
    p_name_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    p_name.paragraph_format.space_after = Pt(2)
    
    # 2. Sub-role Title
    p_role = doc.add_paragraph()
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_role_run = p_role.add_run(role)
    p_role_run.font.size = Pt(12)
    p_role_run.font.italic = True
    p_role_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p_role.paragraph_format.space_after = Pt(4)

    # 3. Contact Info
    contact_parts = [
        f"Phone: {phone}",
        f"Email: {email}",
        f"Location: {location}",
        "GitHub: github.com/prudhviraj",
    ]
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact_run = p_contact.add_run("  |  ".join(contact_parts))
    p_contact_run.font.size = Pt(9.5)
    p_contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p_contact.paragraph_format.space_after = Pt(12)

    def add_section_divider(title):
        p_div = doc.add_paragraph()
        p_div.paragraph_format.space_before = Pt(10)
        p_div.paragraph_format.space_after = Pt(4)
        run = p_div.add_run(title.upper())
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3) # Sleek blue accent
        
        # Add a subtle horizontal bottom border or underline using paragraph styling
        p_div.paragraph_format.keep_with_next = True

    # 4. Professional Summary (Tailored by AI)
    add_section_divider("Professional Summary")
    summary_text = tailored_data.get("summary") or candidate_knowledge.get("professional_summary", "")
    p_sum = doc.add_paragraph()
    p_sum.add_run(summary_text)

    # 5. Core Technical Skills
    add_section_divider("Core Skills")
    skills = candidate_knowledge.get("core_skills", {})
    if skills:
        for category, skill_list in skills.items():
            cat_name = category.replace("_", " ").title()
            p_skill = doc.add_paragraph()
            p_skill.paragraph_format.left_indent = Inches(0.2)
            p_skill.paragraph_format.space_after = Pt(2)
            bold_run = p_skill.add_run(f"{cat_name}: ")
            bold_run.font.bold = True
            p_skill.add_run(", ".join(skill_list))

    # 6. Work Experience (Bullets Tailored by AI)
    work_exp = candidate_knowledge.get("work_experience", [])
    if work_exp:
        add_section_divider("Work Experience")
        for i, exp in enumerate(work_exp):
            if "YOUR COMPANY" in exp.get("company", ""):
                continue
            
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(4)
            p_title.paragraph_format.space_after = Pt(2)
            
            # Left side: Role & Company
            role_run = p_title.add_run(f"{exp.get('role', 'SDE Intern')} — {exp.get('company', 'Company')}")
            role_run.font.bold = True
            
            # Right side: Duration (using spaces to push it right, highly robust for ATS)
            p_title.add_run(f"  ({exp.get('duration', '')})")
            
            # Use AI rewritten bullets if available, else standard achievements
            bullets = []
            if i == 0 and tailored_data.get("experience_bullets"):
                bullets = tailored_data.get("experience_bullets")
            else:
                bullets = exp.get("achievements") or [exp.get("description", "")]
                
            for bullet in bullets:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.left_indent = Inches(0.4)
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(bullet)

    # 7. Key Projects (Targeted highlights)
    projects = candidate_knowledge.get("github_projects", [])
    if projects:
        add_section_divider("Key Projects")
        for proj in projects:
            if "ADD YOUR" in proj.get("name", ""):
                continue
            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_before = Pt(4)
            p_proj.paragraph_format.space_after = Pt(2)
            
            proj_name = p_proj.add_run(proj.get("name", "Project"))
            proj_name.font.bold = True
            
            tech_stack = proj.get("tech_stack", [])
            if tech_stack:
                p_proj.add_run(f"  |  {', '.join(tech_stack)}")
                
            highlights = proj.get("highlights", [])
            for hl in highlights:
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.left_indent = Inches(0.4)
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.add_run(hl)

    # 8. Education
    education = identity.get("education", "B.Tech in Computer Science")
    add_section_divider("Education")
    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.left_indent = Inches(0.2)
    p_edu_run = p_edu.add_run(f"{education}  |  Graduation: {defaults.get('graduation_year', '2024')}")
    
    # Save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info("Custom tailored resume .docx successfully generated at: %s", output_path)
    return output_path

if __name__ == "__main__":
    # Self-test logic
    with open("config/candidate_knowledge.json", "r") as f:
        kb = json.load(f)
    mock_tailored = {
        "summary": "AI Engineer proficient in building RAG systems and LLM pipelines using FastAPI.",
        "experience_bullets": [
            "Engineered high-performance document parsing service handling 50+ matches/sec.",
            "Optimized query performance for vector index matching."
        ]
    }
    build_tailored_resume_docx(kb, mock_tailored, "resumes/test_tailored_resume.docx")
